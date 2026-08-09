"""Deterministic synthetic anexo fixtures (no network, no real tender data)."""

from __future__ import annotations

import io
import zipfile


def make_pdf(pages: int = 1, *, marker: str = "", marker_page: int | None = None) -> bytes:
    """Text PDF; ``marker`` is written only on ``marker_page`` when provided."""
    import fitz  # type: ignore

    document = fitz.open()
    try:
        for index in range(1, pages + 1):
            page = document.new_page()
            text = f"pagina {index}"
            if marker and marker_page == index:
                text = f"{text} {marker}"
            page.insert_text((72, 72), text)
        return document.tobytes()
    finally:
        document.close()


def make_image_only_pdf(pages: int = 1) -> bytes:
    """PDF whose pages carry no extractable text (scanned-document stand-in)."""
    import fitz  # type: ignore

    document = fitz.open()
    try:
        for _ in range(pages):
            page = document.new_page()
            page.draw_rect(fitz.Rect(10, 10, 200, 200), fill=(0, 0, 0))
        return document.tobytes()
    finally:
        document.close()


def make_docx(
    *,
    paragraph: str = "",
    table_rows: list[list[str]] | None = None,
    header: str = "",
    footer: str = "",
    embedded: bool = False,
) -> bytes:
    import docx  # type: ignore

    document = docx.Document()
    if paragraph:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    section = document.sections[0]
    if header:
        section.header.paragraphs[0].text = header
    if footer:
        section.footer.paragraphs[0].text = footer
    buffer = io.BytesIO()
    document.save(buffer)
    payload = buffer.getvalue()
    if embedded:
        payload = _inject_zip_entry(payload, "word/embeddings/oleObject1.bin", b"\x00\x01")
    return payload


def _inject_zip_entry(payload: bytes, name: str, data: bytes) -> bytes:
    source = io.BytesIO(payload)
    target = io.BytesIO()
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(target, "w") as dst:
        for info in src.infolist():
            dst.writestr(info, src.read(info.filename))
        dst.writestr(name, data)
    return target.getvalue()


def make_xlsx(
    sheets: list[str],
    *,
    cells: dict[str, list[tuple[int, int, str]]] | None = None,
    hidden_sheets: tuple[str, ...] = (),
    merged: dict[str, tuple[str, str]] | None = None,
) -> bytes:
    from openpyxl import Workbook  # type: ignore

    workbook = Workbook()
    workbook.active.title = sheets[0]
    for name in sheets[1:]:
        workbook.create_sheet(name)
    for name, entries in (cells or {}).items():
        sheet = workbook[name]
        for row, column, value in entries:
            sheet.cell(row=row, column=column, value=value)
    for name, (span, value) in (merged or {}).items():
        sheet = workbook[name]
        sheet.merge_cells(span)
        sheet[span.split(":")[0]] = value
    for name in hidden_sheets:
        workbook[name].sheet_state = "hidden"
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_csv(rows: int, *, marker: str = "", marker_row: int | None = None) -> bytes:
    lines = ["columna_a,columna_b"]
    for index in range(1, rows + 1):
        if marker and marker_row == index:
            lines.append(f"{marker},{index}")
        else:
            lines.append(f"fila_{index},{index}")
    return ("\n".join(lines)).encode("utf-8")


def make_zip(entries: list[tuple[str, bytes]], *, compress: bool = True) -> bytes:
    buffer = io.BytesIO()
    mode = zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED
    with zipfile.ZipFile(buffer, "w", mode) as archive:
        for name, data in entries:
            archive.writestr(name, data)
    return buffer.getvalue()


def make_zip_bomb(*, member_bytes: int = 8 * 1024 * 1024) -> bytes:
    """Highly compressible single member used to exercise the ratio guard."""
    return make_zip([("bomb.txt", b"A" * member_bytes)], compress=True)


def make_traversal_zip() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("../../escape.txt", b"should not be written")
        archive.writestr("safe.txt", b"contenido seguro")
    return buffer.getvalue()


def _set_encrypted_flag(payload: bytes, target_name: str) -> bytes:
    """Set the general-purpose encryption bit in both headers for one member.

    ``ZipFile.writestr`` rewrites ``flag_bits`` while streaming, so the flag has
    to be patched in the serialized headers instead.
    """
    data = bytearray(payload)
    name = target_name.encode("utf-8")

    def _patch(signature: bytes, flag_offset: int, name_len_offset: int, name_offset: int) -> None:
        cursor = 0
        while True:
            index = data.find(signature, cursor)
            if index == -1:
                return
            cursor = index + 4
            name_length = int.from_bytes(
                data[index + name_len_offset : index + name_len_offset + 2], "little"
            )
            start = index + name_offset
            if data[start : start + name_length] != name:
                continue
            flags = int.from_bytes(
                data[index + flag_offset : index + flag_offset + 2], "little"
            )
            data[index + flag_offset : index + flag_offset + 2] = (flags | 0x1).to_bytes(
                2, "little"
            )

    _patch(b"PK\x03\x04", flag_offset=6, name_len_offset=26, name_offset=30)
    _patch(b"PK\x01\x02", flag_offset=8, name_len_offset=28, name_offset=46)
    return bytes(data)


def make_encrypted_member_zip() -> bytes:
    """Flip the ZIP encryption flag so the member is refused without a password."""
    payload = make_zip(
        [("secreto.txt", b"clasificado"), ("abierto.txt", b"publico")], compress=False
    )
    return _set_encrypted_flag(payload, "secreto.txt")


OLE_HEADER = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def make_fake_legacy_ole(extra: bytes = b"\x00" * 512) -> bytes:
    return OLE_HEADER + extra
